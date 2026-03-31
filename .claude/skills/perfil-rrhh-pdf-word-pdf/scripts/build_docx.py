#!/usr/bin/env python3
"""Build a .docx from JSON with title, optional HR analysis block, sections, and professional color themes."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

# Paletas profesionales: A = corporativo (azul institucional), B = ejecutivo (teal/gris sobrio)
THEMES: dict[str, dict[str, RGBColor]] = {
    "corporate": {
        "title": RGBColor(0x1A, 0x36, 0x5F),
        "heading": RGBColor(0x2C, 0x52, 0x7F),
        "body_muted": RGBColor(0x33, 0x33, 0x33),
        "accent": RGBColor(0x44, 0x72, 0xC4),
    },
    "executive": {
        "title": RGBColor(0x1C, 0x28, 0x2C),
        "heading": RGBColor(0x00, 0x5F, 0x6B),
        "body_muted": RGBColor(0x33, 0x33, 0x33),
        "accent": RGBColor(0x00, 0x96, 0x88),
    },
}


def _apply_line_spacing(p) -> None:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.18


def _style_title(paragraph, text: str, colors: dict[str, RGBColor]) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = colors["title"]
    run.font.name = "Calibri"
    _apply_line_spacing(paragraph)


def _add_section_heading(doc: Document, text: str, colors: dict[str, RGBColor], level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    _apply_line_spacing(p)
    for run in p.runs:
        run.font.color.rgb = colors["heading"]
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(14)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(12)
            run.bold = True


def add_body_paragraphs(doc: Document, body: str, colors: dict[str, RGBColor]) -> None:
    for block in body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        for line in block.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith(("- ", "* ", "• ")):
                p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                p = doc.add_paragraph(line)
            _apply_line_spacing(p)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"
                run.font.color.rgb = colors["body_muted"]


def add_hr_analysis_block(doc: Document, heading: str, body: str, colors: dict[str, RGBColor]) -> None:
    """Bloque destacado: título nivel 2 + cuerpo con sangría (análisis RR.HH. sustancial)."""
    if not body.strip():
        return
    _add_section_heading(doc, heading, colors, level=2)
    for block in body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        for line in block.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith(("- ", "* ", "• ")):
                p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(0.2)
            _apply_line_spacing(p)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"
                run.font.color.rgb = colors["body_muted"]
    doc.add_paragraph()


def build_docx(payload: dict[str, Any], out_path: Path) -> None:
    theme_key = (payload.get("theme") or "corporate").lower()
    if theme_key not in THEMES:
        theme_key = "corporate"
    colors = THEMES[theme_key]

    title = payload.get("title") or "Perfil profesional"
    sections = payload.get("sections")
    if not isinstance(sections, list):
        raise ValueError('"sections" debe ser una lista de {heading, body}.')

    doc = Document()
    sect = doc.sections[0]
    sect.left_margin = Inches(1)
    sect.right_margin = Inches(1)

    title_p = doc.add_paragraph()
    _style_title(title_p, title, colors)

    hr = payload.get("hr_analysis")
    if isinstance(hr, dict):
        h = (hr.get("heading") or "Análisis RR.HH. — Diagnóstico del perfil").strip()
        b = (hr.get("body") or "").strip()
        if b:
            add_hr_analysis_block(doc, h, b, colors)

    for item in sections:
        if not isinstance(item, dict):
            continue
        heading = (item.get("heading") or "").strip()
        body = (item.get("body") or "").strip()
        if heading:
            _add_section_heading(doc, heading, colors, level=1)
        if body:
            add_body_paragraphs(doc, body, colors)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Genera .docx desde JSON (temas corporate | executive).")
    parser.add_argument("--json", "-j", type=Path, required=True, help="Archivo JSON")
    parser.add_argument("--out", "-o", type=Path, required=True, help="Salida .docx")
    args = parser.parse_args()
    if not args.json.is_file():
        print(f"Error: no existe: {args.json}", file=sys.stderr)
        return 1
    try:
        payload = json.loads(args.json.read_text(encoding="utf-8"))
        if not isinstance(payload, dict):
            raise ValueError("El JSON raíz debe ser un objeto.")
        build_docx(payload, args.out)
    except (json.JSONDecodeError, ValueError, KeyError) as e:
        print(f"Error en JSON o contenido: {e}", file=sys.stderr)
        return 1
    except Exception as e:
        print(f"Error al generar docx: {e}", file=sys.stderr)
        return 1
    print(f"Generado: {args.out}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
